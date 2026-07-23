# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/claude_user/test/remote_backup/NetBackup_CUBRID_정책구성가이드.docx"
doc = Document()
st = doc.styles['Normal']; st.font.name='맑은 고딕'
st.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.font.size=Pt(10.5)
BLUE=RGBColor(0x1F,0x4E,0x79); GRAY=RGBColor(0x59,0x59,0x59); RED=RGBColor(0xB0,0x2A,0x2A)
def kf(r,name='맑은 고딕'):
    r.font.name=name; rPr=r._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'),name)
def shade(cell,c):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),c); cell._tc.get_or_add_tcPr().append(sh)
def title(t,s):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=BLUE; kf(r)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run(s); r2.font.size=Pt(10); r2.font.color.rgb=GRAY; kf(r2); doc.add_paragraph()
def h(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(13.5); r.font.color.rgb=BLUE; kf(r)
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bo=OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6'); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),'1F4E79')
    b.append(bo); pPr.append(b)
def h2(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=RGBColor(0x2E,0x5A,0x88); kf(r)
def para(t,bold=False,color=None,size=10.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; kf(r); r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t,bp=None):
    p=doc.add_paragraph(style='List Bullet')
    if bp: r=p.add_run(bp); r.bold=True; kf(r); r.font.size=Pt(10.5)
    r2=p.add_run(t); kf(r2); r2.font.size=Pt(10.5)
def code(lines):
    for ln in lines.split("\n"):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
        pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F2F2'); pPr.append(sh)
        r=p.add_run(ln if ln else " "); r.font.size=Pt(9)
        rp=r._element.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Consolas'); rf.set(qn('w:hAnsi'),'Consolas'); rp.append(rf)
def table(headers,rows,widths=None,fs=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=t.rows[0].cells
    for i,x in enumerate(headers):
        hc[i].text=''; r=hc[i].paragraphs[0].add_run(x); r.bold=True; r.font.size=Pt(fs); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); kf(r); shade(hc[i],'1F4E79')
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].text=''; r=c[i].paragraphs[0].add_run(v); r.font.size=Pt(fs); kf(r)
    if widths:
        for rr in t.rows:
            for i,w in enumerate(widths): rr.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def pb(): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

title("CUBRID + Veritas NetBackup 정책 구성 가이드",
      "관리자용  |  CUBRID 데이터베이스를 NetBackup 으로 백업하기 위한 정책/스케줄 구성  |  2026-07")

h("1. 개요")
para("CUBRID 는 NetBackup 전용 데이터베이스 에이전트가 없다. 따라서 'CUBRID 백업파일을 만든 뒤 NetBackup 이 그 파일을 수집(파일 기반)'하는 방식으로 연동한다. 연동 방법은 두 가지이며, 본 문서는 NetBackup 관리자가 준비해야 할 정책 구성을 설명한다.")
table(["방법","개시 주체","NetBackup 구성 요점"],
 [["방법 1: 스크립트 기반(권장)","DB 서버의 cron 스크립트","User Backup 스케줄을 가진 Standard 정책 + 클라이언트 등록"],
  ["방법 2: 정책 기반","NetBackup 스케줄","Standard 정책의 Backup Selections=스테이징 폴더 + 사전 스크립트(bpstart_notify)"]],
 widths=[2.3,1.9,2.4])

h("2. 사전 준비")
bullet("DB 서버를 NetBackup 클라이언트로 설치하고 마스터 서버에 클라이언트로 등록한다.")
bullet("마스터-클라이언트 연결 확인: 클라이언트에서 bptestbpcd, 또는 bpclntcmd -pn / bpclntcmd -hn <master>.")
bullet("스토리지 유닛(디스크/테이프/dedup)과 보존 레벨을 준비한다.")
bullet("DB 서버에 백업 스테이징 공간을 확보한다(백업 크기 + 여유). 예: /backup/cubrid_stage/<DB>.")
bullet("클라이언트/마스터 시간 동기(NTP). 대용량이면 timeout 여유를 둔다.")

pb()
h("3. 방법 1 : 스크립트 기반 (권장)")
para("DB 서버에서 제공 스크립트(netbackup_cubrid_backup.sh)를 cron 으로 실행한다. 스크립트가 cubrid backupdb 로 스테이징 백업 후, bpbackup(사용자 개시 백업)으로 NetBackup 에 수집시킨다.")
h2("3.1 NetBackup 정책 구성(관리자)")
table(["항목","설정 값(예시)","설명"],
 [["Policy type","Standard","UNIX/Linux 파일 시스템 백업"],
  ["Policy name","CUBRID_FS","스크립트의 POLICY 와 일치"],
  ["Clients","<DB 서버 호스트명>","백업 대상 클라이언트"],
  ["Storage unit","(사이트 스토리지)","디스크/테이프/dedup"],
  ["Schedule (type)","User Backup","bpbackup 사용자 개시 허용 스케줄. 스크립트의 SCHEDULE 과 일치"],
  ["Retention","(정책)","백업 보존 기간"],
  ["Keyword","cubrid_<DB>_<날짜>","스크립트 KEYWORD (검색/식별용, 선택)"]],
 widths=[1.7,2.3,2.6])
h2("3.2 DB 서버(클라이언트) 실행 - cron 예시")
code("# 매일 02:00 전체 백업\n0 2 * * *  DB=demodb LEVEL=0 STAGE=/backup/cubrid_stage/demodb \\\n           POLICY=CUBRID_FS SCHEDULE=Default-Application-Backup \\\n           KEEP_STAGE=no bash /path/netbackup_cubrid_backup.sh")
para("스크립트는 bpbackup 실패 시 재시도(기본 3회)를 포함한다. 정책명/스케줄명은 위 3.1 구성과 반드시 일치시킨다.", color=GRAY, size=9.5)

h("4. 방법 2 : 정책 기반 (NetBackup 스케줄이 개시)")
para("NetBackup 스케줄이 백업을 개시하고, 정책의 사전 스크립트에서 CUBRID 백업파일을 미리 생성한다.")
h2("4.1 NetBackup 정책 구성")
table(["항목","설정 값(예시)","설명"],
 [["Policy type","Standard","파일 시스템 백업"],
  ["Backup Selections","/backup/cubrid_stage/demodb","CUBRID 백업파일이 놓이는 스테이징 폴더"],
  ["Schedule (type)","Full / Differential","정기 스케줄(빈도/보존)"],
  ["Clients","<DB 서버 호스트명>",""],
  ["Storage unit / Retention","(사이트)",""]],
 widths=[1.9,2.3,2.4])
h2("4.2 사전/사후 스크립트(클라이언트에 배치)")
bullet("bpstart_notify.<policy> : 백업 시작 전에 실행. 여기서 cubrid backupdb 로 스테이징에 백업 생성.")
bullet("bpend_notify.<policy> : 백업 종료 후 실행. 스테이징 정리 등.")
code("# 위치: /usr/openv/netbackup/bin/bpstart_notify.CUBRID_FS (실행권한 필요)\n#!/bin/bash\nexport CUBRID=/home/cubrid/CUBRID; export PATH=$CUBRID/bin:$PATH\nexport CUBRID_DATABASES=$CUBRID/databases; export LD_LIBRARY_PATH=$CUBRID/lib:$CUBRID/cci/lib\nrm -rf /backup/cubrid_stage/demodb; mkdir -p /backup/cubrid_stage/demodb\ncubrid backupdb -D /backup/cubrid_stage/demodb -l 0 --no-check demodb")
para("주의: bpstart_notify 는 백업 시작을 지연시키므로, 백업 생성 시간을 감안해 정책의 시작 timeout(예 client read timeout / BPSTART_TIMEOUT)을 충분히 크게 둔다.", color=GRAY, size=9.5)

pb()
h("5. 정책 구성 항목 요약")
table(["구성 항목","권장/예시"],
 [["Policy type","Standard(UNIX/Linux 파일)"],
  ["방법 1 스케줄","User Backup 유형(스크립트 bpbackup 개시)"],
  ["방법 2 스케줄","Full/Differential + bpstart_notify 사전 스크립트"],
  ["Backup Selection","CUBRID 스테이징 폴더(/backup/cubrid_stage/<DB>)"],
  ["Storage unit","디스크/dedup 권장(대용량·중복제거)"],
  ["Retention","백업 정책에 따라(예: 전체 4주, 증분 1주)"],
  ["압축/중복제거","스토리지 유닛/어플라이언스(dedup) 기능 활용(스크립트 무관)"]],
 widths=[2.4,4.2])

h("6. 검증 및 복원 (관리자)")
para("NetBackup 에 수집된 CUBRID 백업파일을 확인/복원한 뒤, DBA 가 CUBRID 복원을 수행한다.")
h2("6.1 백업 이미지 목록")
code("bplist -C <client> -t 0 -R /backup/cubrid_stage")
h2("6.2 파일 복원")
code("bprestore -C <client> -D <client> -t 0 -w -L /tmp/rest.log /backup/cubrid_stage/demodb")
h2("6.3 CUBRID 복원 (DBA)")
code("cubrid restoredb -B <복원된_스테이징_경로> demodb\ncubrid checkdb --SA-mode demodb")
para("정기적으로 복원 리허설(restore + checkdb)을 수행해 백업 유효성을 확인할 것.", color=GRAY, size=9.5)

h("7. 운영 권고 및 주의사항")
bullet("스테이징 공간은 백업 크기 이상 확보(전체 백업은 DB 크기에 준함). 수집 후 자동 정리(KEEP_STAGE=no) 권장.")
bullet("백업 레벨 매핑: 전체(Level 0)는 주기적으로, 증분(Level 1/2)은 그 사이에. 스케줄로 분리 운영.")
bullet("대용량 백업은 bpbackup -w 대기시간·정책 timeout·클라이언트 read timeout 을 충분히 설정.")
bullet("bpbackup 플래그·스케줄 유형은 NetBackup 버전/정책 유형에 따라 다를 수 있으므로 사이트 기준으로 검증.")
bullet("로그: NetBackup 클라이언트 로그(/usr/openv/netbackup/logs/bpbackup 등)와 스크립트 로그를 함께 보관.")
bullet("정책명/스케줄명/클라이언트명은 DBA(스크립트 설정)와 NetBackup 관리자가 사전에 합의해 일치시킬 것.")

doc.save(OUT); print("saved:", OUT)
