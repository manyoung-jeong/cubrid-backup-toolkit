# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/claude_user/test/remote_backup/원격백업_사용자매뉴얼.docx"
doc = Document()
st = doc.styles['Normal']; st.font.name='맑은 고딕'
st.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.font.size=Pt(10.5)
BLUE=RGBColor(0x1F,0x4E,0x79); GRAY=RGBColor(0x59,0x59,0x59); RED=RGBColor(0xB0,0x2A,0x2A); GREEN=RGBColor(0x2E,0x6B,0x2E)

def kf(r,name='맑은 고딕'):
    r.font.name=name; rPr=r._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'),name)
def shade(cell,c):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),c); cell._tc.get_or_add_tcPr().append(sh)
def title(t,s):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=BLUE; kf(r)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run(s); r2.font.size=Pt(10.5); r2.font.color.rgb=GRAY; kf(r2); doc.add_paragraph()
def h(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=BLUE; kf(r)
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bo=OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6'); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),'1F4E79')
    b.append(bo); pPr.append(b)
def h2(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=RGBColor(0x2E,0x5A,0x88); kf(r)
def step(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(1)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79); kf(r)
def para(t,bold=False,color=None,size=10.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; kf(r); r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def kv(label,text,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label); r.bold=True; kf(r); r.font.size=Pt(10.5)
    if color: r.font.color.rgb=color
    r2=p.add_run(text); kf(r2); r2.font.size=Pt(10.5)
def bullet(t,bp=None):
    p=doc.add_paragraph(style='List Bullet')
    if bp: r=p.add_run(bp); r.bold=True; kf(r); r.font.size=Pt(10.5)
    r2=p.add_run(t); kf(r2); r2.font.size=Pt(10.5)
def num(t):
    p=doc.add_paragraph(style='List Number'); r=p.add_run(t); kf(r); r.font.size=Pt(10.5)
def code(lines):
    for ln in lines.split("\n"):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
        pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F2F2'); pPr.append(sh)
        r=p.add_run(ln if ln else " "); r.font.size=Pt(9)
        rp=r._element.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Consolas'); rf.set(qn('w:hAnsi'),'Consolas'); rp.append(rf)
def table(headers,rows,widths=None,fs=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=t.rows[0].cells
    for i,x in enumerate(headers):
        hc[i].text=''; r=hc[i].paragraphs[0].add_run(x); r.bold=True; r.font.size=Pt(fs); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); kf(r); shade(hc[i],'1F4E79')
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].text=''; r=c[i].paragraphs[0].add_run(v); r.font.size=Pt(fs); kf(r)
    if widths:
        for rr in t.rows:
            for i,w in enumerate(widths): rr.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def pb(): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

title("CUBRID 원격 백업 인터페이스 사용자 매뉴얼",
      "cubrid backupdb 결과를 다른 IP 의 백업 서버로 전송  |  대상: DBA/운영자  |  2026-07")

# ---- 1. 개요 ----
h("1. 개요")
para("cubrid backupdb 의 백업 결과를 다른 IP 를 가진 백업 서버로 전송하는 도구이다. 두 대의 서버를 사용한다.")
bullet("백업을 뜨는 서버(운영 DB 가 있는 곳).", "DB 서버(송신) : ")
bullet("백업 파일을 저장하는 다른 IP 의 장비.", "백업 서버(수신) : ")
para("아래 두 가지 조건을 만족하도록 만들어졌다.")
num("네트워크가 가끔 끊겨도 백업이 정상 완료된다(재접속/이어받기).")
num("백업 장비가 지정 시간(기본 60초) 이상 문제면 오류 로그를 남기고 backupdb 를 정지한다.")

h2("어느 방식을 쓸까? (한눈에)")
table(["상황","권장 방식"],
 [["대용량이거나 네트워크가 자주 끊김","방식 A (재개형 전송) - 이어받기 지원"],
  ["로컬 디스크 공간이 부족 / 빠르고 안정적인 망","방식 B (ssh 스트리밍) - 로컬 공간 불필요"],
  ["운영 영향을 최소화(커밋 지연 절대 회피)","방식 C (로컬 백업 후 전송)"]],
 widths=[3.4,3.2])

pb()
# ---- 2. 빠른 시작 ----
h("2. 빠른 시작 : 서버 설정과 백업 실행 (단계별)")
para("아래 순서대로 따라 하면 된다. 예시는 백업 서버 IP 를 192.168.7.39, 포트 9099, DB 를 demodb 로 가정한다.")

h2("[방식 A] 재개형 전송 (권장)")
step("1단계. 파일 배치")
bullet("백업 서버(수신)에 rbk_listener.c 를, DB 서버(송신)에 remote_backupdb.sh 와 rbk_forward.c 를 둔다.")
step("2단계. 백업 서버(수신) 설정 - 리스너 준비 후 대기")
code("gcc -O2 -o rbk_listener rbk_listener.c        # 최초 1회 빌드\nmkdir -p /root/db_backup                        # 받을 폴더\n./rbk_listener 9099 /root/db_backup/demodb_$(date +%Y%m%d).bk /root/db_backup/rbk.log")
para("'listener 시작 ...' 로그가 뜨면 수신 대기 상태이다. (백업이 끝나면 자동 종료)", color=GRAY, size=9.5)
step("3단계. DB 서버(송신) 설정 - CUBRID 환경변수 확인")
code("echo $CUBRID        # 비어 있으면 아래처럼 설정(경로는 환경에 맞게)\nexport CUBRID=/home/xxx/CUBRID\nexport CUBRID_DATABASES=$CUBRID/databases\nexport PATH=$CUBRID/bin:$PATH\nexport LD_LIBRARY_PATH=$CUBRID/lib:$CUBRID/cci/lib")
step("4단계. DB 서버(송신)에서 백업 실행")
code("REMOTE_IP=192.168.7.39 REMOTE_PORT=9099 DB=demodb LEVEL=0 TIMEOUT=60 \\\n  bash remote_backupdb.sh")
para("(rbk_forward 는 스크립트가 자동으로 빌드한다.)", color=GRAY, size=9.5)
step("5단계. 결과 확인")
bullet("로그에 '백업 및 원격 전송 완료' -> 성공. 백업 서버의 파일 크기를 확인한다.")
bullet("로그에 'backupdb 정지' -> 백업 장비 문제(조건2). 네트워크/수신 서버 점검 후 재실행.")

h2("[방식 B] ssh 스트리밍 직접 전송 (로컬 공간 불필요)")
step("1단계. SSH 키 1회 등록 (이후 비밀번호 불필요)")
code("ssh-keygen -t ed25519 -N '' -f ~/.ssh/bk_key\nssh-copy-id -i ~/.ssh/bk_key.pub root@192.168.7.39   # 최초 1회만 비밀번호 입력")
step("2단계. 백업 서버에 받을 폴더 생성")
code('ssh -i ~/.ssh/bk_key root@192.168.7.39 "mkdir -p /root/db_backup"')
step("3단계. DB 서버에서 스트리밍 백업 실행")
code('FIFO=/tmp/bkfifo; rm -f $FIFO; mkfifo $FIFO\ncat $FIFO | ssh -i ~/.ssh/bk_key root@192.168.7.39 \\\n     "cat > /root/db_backup/demodb_$(date +%Y%m%d_%H%M%S).bk" &\ncubrid backupdb -D $FIFO -l 0 --no-check demodb\nrm -f $FIFO')
para("(예시 스크립트 stream_backup_nhidb.sh 의 대상 DB/경로만 바꿔 재사용해도 된다.)", color=GRAY, size=9.5)
step("4단계. 결과 확인")
code('ssh -i ~/.ssh/bk_key root@192.168.7.39 "ls -lh /root/db_backup/"')

pb()
# ---- 3. 구성 파일 ----
h("3. 구성 파일")
table(["파일","실행 위치","역할"],
 [["remote_backupdb.sh","DB 서버(송신)","방식 A 메인 실행 스크립트"],
  ["rbk_forward.c / rbk_forward","DB 서버(송신)","FIFO->로컬 스풀 흡수 후 원격 전송. 이어받기 + 무응답 watchdog"],
  ["rbk_listener.c / rbk_listener","백업 서버(수신)","재접속/이어받기 지원 수신기"],
  ["stream_backup_nhidb.sh","DB 서버(송신)","방식 B 예시(ssh 스트리밍) 스크립트"],
  ["e2e_test.sh","DB 서버","백업->전송->복원->검증 종단 시험"]],
 widths=[2.2,1.6,3.0])

# ---- 4. 사전 요구사항 ----
h("4. 사전 요구사항")
bullet("송신측: CUBRID 환경변수, 대상 DB 준비, gcc(포워더 빌드).")
bullet("수신측: gcc(리스너 빌드) 또는 빌드된 rbk_listener. 방식 B 는 ssh 접근만 필요.")
bullet("네트워크: 방식 A 는 전송 포트(기본 9099) 개방. 방식 B 는 22(ssh)만.")
bullet("디스크: 방식 A 는 송신측에 백업 크기만큼 스풀 공간 필요. 방식 B 는 거의 불필요.")

# ---- 5. 방식 A 상세 ----
h("5. 방식 A 상세 (환경변수/동작/반환코드)")
h2("5.1 환경변수")
table(["변수","기본값","설명"],
 [["DB","demodb","백업 대상 DB"],
  ["LEVEL","0","백업 레벨 0(전체)/1/2"],
  ["REMOTE_IP","(필수)","백업 서버 IP"],
  ["REMOTE_PORT","9099","리스너 포트"],
  ["TIMEOUT","60","무응답 허용 시간(초). 초과 시 backupdb 정지"],
  ["SPOOL","작업폴더/spool_<DB>","로컬 스풀 파일(백업 크기만큼 공간 필요)"],
  ["LOG","작업폴더/remote_backup.log","로그 파일"],
  ["CUBRID_BK_OPT","--no-check","backupdb 추가 옵션(예: --SA-mode)"]],
 widths=[1.6,2.1,3.1])
h2("5.2 동작")
bullet("포워더가 FIFO 를 로컬 스풀로 빠르게 흡수 -> backupdb 가 네트워크 때문에 멈추지 않음(커밋 지연 회피).")
bullet("순단 시 재접속하여 수신측 오프셋부터 이어받음(조건1). 끝에 총 크기 일치 검증.")
bullet("원격이 TIMEOUT 초 이상 무응답이면 오류 로그 후 종료(반환 2) -> backupdb 정지(조건2).")
h2("5.3 반환코드")
table(["코드","의미"],
 [["0","백업 및 원격 전송 완료"],["2","백업 장비 무응답 -> backupdb 정지(조건2)"],["1","기타 오류"]],
 widths=[1.0,5.6])

# ---- 6. 방식 B 특징 ----
h("6. 방식 B 특징 / 주의")
bullet("장점: 로컬 백업 사본 불필요(디스크 절약), 백업과 동시에 실시간 전송. FIFO 는 데이터를 저장하지 않는 임시 파이프(크기 0)이다.")
bullet("주의1(재개 불가): 전송 중 끊기면 처음부터 다시.")
bullet("주의2(커밋 지연): 네트워크가 느려 파이프가 막히면 운영 DB 커밋이 지연될 수 있음. 안정망에서만 권장.")

# ---- 7. 방식 C ----
h("7. (참고) 방식 C : 로컬 백업 후 전송 (운영 안전 최우선)")
code("cubrid backupdb -D /local/bk --no-check demodb\nrsync -a --partial --append-verify --timeout=20 /local/bk/ root@192.168.7.39:/root/db_backup/")

# ---- 8. 백업 장비 유형별 연동 방법 ----
pb()
h("8. 백업 장비 유형별 연동 방법")
para("백업 대상 장비가 무엇을 지원하느냐에 따라 연동 방법이 다르다. 폐쇄형 어플라이언스는 우리 커스텀 리스너(방식 A)를 못 올리므로, 표준 프로토콜이나 장비의 정식 수집 경로로 연동한다.")
table(["백업 장비 유형","노출 프로토콜","CUBRID 연동 방법","이어받기"],
 [["일반 리눅스 서버","SSH","방식 B(ssh 스트리밍) 또는 방식 A(리스너)","방식 A: 가능"],
  ["리눅스 NAS(프로그램 실행 가능)","SSH/자체","방식 A(리스너 설치)","가능"],
  ["NFS/CIFS 공유 NAS","NFS/SMB 마운트","공유 마운트 후 cubrid backupdb -D /mnt/... 직접","마운트 재연결 의존"],
  ["S3/오브젝트 스토리지","S3 API","로컬 백업 후 aws s3 cp, 또는 cat FIFO | aws s3 cp - s3://..","aws cli 멀티파트 재시도"],
  ["Veritas NetBackup","NBU 클라이언트/bpbackup","스테이징 백업 후 bpbackup 수집(전용 스크립트)","NetBackup 이 재시도 관리"],
  ["테이프/상용 백업SW","백업SW 경유","백업SW 정책이 스테이징 파일 수집","백업SW 관리"],
  ["Data Domain 등 dedup","NFS/CIFS/OST/DDBoost","마운트 직접 또는 백업SW 경유","프로토콜/SW 의존"]],
 widths=[1.7,1.5,2.5,1.0])
h2("8.1 Veritas NetBackup 연동")
para("NetBackup 은 CUBRID 전용 에이전트가 없으므로, 'CUBRID 백업파일을 만든 뒤 NetBackup 이 그 파일을 수집'하는 파일 기반 연동이 표준이다. 두 가지 방법이 있다.")
bullet("스크립트가 cubrid backupdb 로 스테이징 백업 후 bpbackup 으로 즉시 NBU 에 수집. 제공 스크립트: netbackup_cubrid_backup.sh", "스크립트 기반(권장, 자족적) : ")
bullet("NetBackup 정책의 백업 대상에 스테이징 폴더를 지정하고, 정책 사전 스크립트(bpstart_notify)에서 cubrid backupdb 를 실행.", "정책 기반 : ")
para("실행 예시(스크립트 기반):")
code('DB=demodb LEVEL=0 STAGE=/backup/cubrid_stage/demodb \\\n  POLICY=CUBRID_FS SCHEDULE=Default-Application-Backup \\\n  bash netbackup_cubrid_backup.sh')
para("사전 요구: 이 호스트가 NetBackup 클라이언트이고 마스터에 정책/스케줄이 구성돼야 한다. bpbackup 기본 경로는 /usr/openv/netbackup/bin. 정책명/스케줄/키워드는 NetBackup 관리자에게 확인한다(플래그는 NBU 버전·정책 유형에 따라 다를 수 있음). 스크립트는 네트워크/장비 순단 대비 bpbackup 재시도(기본 3회)를 포함한다.", color=GRAY, size=9.5)

# ---- 9. 검증 ----
h("9. 백업본 검증(복원)")
para("원격 백업본이 실제 복원되는지 임시 위치에서 확인(운영 DB 와 무관).")
code("cubrid restoredb -B <백업디렉토리> <DB>\ncubrid checkdb --SA-mode <DB>")
para("e2e_test.sh 는 격리 DB 로 백업->전송->restoredb->checkdb->행수 일치까지 자동 검증한다.")

# ---- 9. 실제 사례 ----
h("10. 실제 적용 사례 (검증됨)")
table(["항목","값"],
 [["대상","운영 중 nhidb (약 168GB, Full/Level 0 온라인 백업)"],
  ["방식","방식 B (ssh 스트리밍 직접 전송)"],
  ["대상 파일","root@192.168.7.39:/root/db_backup/nhidb_20260723_103908.bk"],
  ["전송 크기","약 35GB (LZ4 압축)"],
  ["소요 시간","약 5분"],
  ["결과","backupdb rc=0, ssh rc=0, 정상 완료. nhidb 운영 지속"]],
 widths=[1.6,5.0])

# ---- 10. 트러블슈팅 ----
h("11. 트러블슈팅")
table(["증상","원인/조치"],
 [["반환 2 로 종료","백업 장비 무응답. 원격 서버/네트워크/포트 확인 후 재실행"],
  ["연결 실패","포트(방식A)/ssh(방식B) 접근 불가. 방화벽/IP/키 확인"],
  ["스풀 공간 부족(방식A)","SPOOL 디스크 여유 확보 또는 방식 B 사용"],
  ["운영 DB 커밋 지연(방식B)","네트워크 지연. 방식 A(스풀) 또는 방식 C(로컬) 사용"],
  ["restore/checkdb 실패","checkdb 는 서버 정지 상태 --SA-mode 로 실행. 백업파일 이름/경로 확인"]],
 widths=[2.2,4.4])

# ---- 11. 주의/한계 ----
h("12. 주의사항 및 한계")
bullet("방식 B(ssh 스트리밍)는 전송 중 끊기면 재개되지 않는다. 대용량/불안정망은 방식 A 권장.")
bullet("방식 A 는 송신측 로컬 스풀 공간이 백업 크기만큼 필요하다.")
bullet("온라인 백업 스트리밍이 네트워크에 막히면 운영 DB 커밋이 지연될 수 있다(방식 A 스풀로 완화, 방식 C 완전 회피).")
bullet("전송 완료 후 백업본은 정기적으로 복원 검증(restoredb+checkdb) 권장.")
bullet("방식 B 용 SSH 키는 필요 없어지면 원격 authorized_keys 와 로컬 키를 제거한다.")

doc.save(OUT); print("saved:", OUT)
