# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/claude_user/test/remote_backup/실환경_백업검증_체크리스트.docx"
doc = Document()
st = doc.styles['Normal']; st.font.name='맑은 고딕'
st.element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.font.size=Pt(10.5)
BLUE=RGBColor(0x1F,0x4E,0x79); GRAY=RGBColor(0x59,0x59,0x59); RED=RGBColor(0xB0,0x2A,0x2A)
def kf(r,name='맑은 고딕'):
    r.font.name=name; rPr=r._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'),name)
def shade(cell,c):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),c); cell._tc.get_or_add_tcPr().append(sh)
def title(t,s):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=BLUE; kf(r)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run(s); r2.font.size=Pt(10); r2.font.color.rgb=GRAY; kf(r2); doc.add_paragraph()
def h(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(13.5); r.font.color.rgb=BLUE; kf(r)
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bo=OxmlElement('w:bottom')
    bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6'); bo.set(qn('w:space'),'3'); bo.set(qn('w:color'),'1F4E79')
    b.append(bo); pPr.append(b)
def h2(t):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=RGBColor(0x2E,0x5A,0x88); kf(r)
def para(t,bold=False,color=None,size=10.5):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; kf(r); r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def chk(t):  # 체크박스 항목
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15); p.paragraph_format.space_after=Pt(1)
    r=p.add_run("[  ] "); r.bold=True; kf(r); r.font.size=Pt(10.5)
    r2=p.add_run(t); kf(r2); r2.font.size=Pt(10.5)
def bullet(t,bp=None):
    p=doc.add_paragraph(style='List Bullet')
    if bp: r=p.add_run(bp); r.bold=True; kf(r); r.font.size=Pt(10.5)
    r2=p.add_run(t); kf(r2); r2.font.size=Pt(10.5)
def code(lines):
    for ln in lines.split("\n"):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
        pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F2F2'); pPr.append(sh)
        r=p.add_run(ln if ln else " "); r.font.size=Pt(9)
        rp=r._element.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Consolas'); rf.set(qn('w:hAnsi'),'Consolas'); rp.append(rf)
def table(headers,rows,widths=None,fs=8.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=t.rows[0].cells
    for i,x in enumerate(headers):
        hc[i].text=''; r=hc[i].paragraphs[0].add_run(x); r.bold=True; r.font.size=Pt(fs); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); kf(r); shade(hc[i],'1F4E79')
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].text=''; r=c[i].paragraphs[0].add_run(v); r.font.size=Pt(fs); kf(r)
    if widths:
        for rr in t.rows:
            for i,w in enumerate(widths): rr.cells[i].width=Inches(w)
    doc.add_paragraph(); return t
def pb(): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

title("CUBRID 장비별 백업 실환경 검증 체크리스트",
      "S3 / NFS(CIFS) NAS / NetBackup - mock 없이 실제 대상에서 검증  |  대상: DBA/운영자  |  2026-07")

h("1. 개요")
para("격리(mock) 종단 시험(e2e_devices_test.sh)을 통과한 백업 스크립트를 실제 대상 장비에서 검증할 때 사용하는 체크리스트다. 각 방식마다 사전조건 -> 실행 명령 -> 기대결과 -> 복원 검증 순으로 확인한다.")
para("주의: 복원 검증은 반드시 운영 DB 가 아닌 별도 DB 이름/호스트(스크래치)에서 수행한다. restoredb 는 대상 DB 를 덮어쓴다.", bold=True, color=RED)

h("2. 공통 사전조건")
chk("송신(DB) 서버에 CUBRID 환경변수 설정 확인(echo $CUBRID)")
chk("백업 대상 DB 준비 및 상태 정상")
chk("복원 검증용 여유 공간과 스크래치 DB 이름 확보(운영 DB 와 분리)")
chk("운영 영향 고려(온라인 백업 시간대, 네트워크 대역폭)")
chk("방식 A(재개형) 사용 시 gcc 로 rbk_forward/rbk_listener 빌드")

h("3. 공통 검증(복원) 절차")
para("각 방식에서 대상에 저장된 백업본을 임시 위치로 가져온 뒤 아래로 검증한다.")
code("# 백업본이 있는 디렉토리에 <DB>_bk0v000 형태 파일이 있어야 함\ncubrid restoredb -B <백업본_디렉토리> <스크래치DB>\ncubrid checkdb --SA-mode <스크래치DB>\n# (선택) 행수/데이터 확인\ncubrid server start <스크래치DB> && csql -u dba <스크래치DB> -c \"SELECT count(*) FROM <표>\"")
para("통과 기준(공통): 백업 스크립트 rc=0 + 대상에 백업파일 존재(크기>0) + restoredb rc=0 + checkdb rc=0 (+ 가능하면 원본과 행수 일치).", bold=True)

h2("3.1 자동 검증 스크립트 (real-run) - realrun_verify.sh")
para("위 절차를 자동화한다. 백업 실행 -> 대상에서 백업본 취득 -> 격리 위치로 복원(restoredb -u -p) -> checkdb -> PASS/FAIL 출력.")
bullet("안전: 별도 CUBRID_DATABASES + 격리 경로로 복원하므로 운영 DB 파일을 절대 덮어쓰지 않는다(-u 옵션).")
bullet("가드: 이 호스트에서 동일 이름 DB 서버가 구동 중이면 충돌 방지를 위해 중단한다(검증 전용 호스트 권장, RUN_ANYWAY=yes 로 강제 가능하나 비권장).")
para("방식별 실행:")
code("# 이미 만들어진 백업 디렉토리 검증(대상 무관)\nMETHOD=dir BK_DIR=<백업디렉토리> DB=<db> bash realrun_verify.sh\n# S3: 업로드 후 다운로드하여 검증\nMETHOD=s3 DB=<db> S3_URI=s3://<bucket>/cubrid/<db> [AWS_PROFILE=..] [AWS_ENDPOINT=..] bash realrun_verify.sh\n# NFS: NAS 백업 후 검증\nMETHOD=nfs DB=<db> MOUNT_SRC=<nas>:/vol MOUNTPOINT=/mnt/nasbackup DEST_SUBDIR=cubrid/<db> bash realrun_verify.sh\n# NetBackup: 백업 + bprestore 후 검증\nMETHOD=nbu DB=<db> POLICY=<정책> SCHEDULE=<스케줄> CLIENT=<client> bash realrun_verify.sh")
para("판정: 종료코드 0=PASS, 1=FAIL. 로그(realrun_verify.log)에 restoredb/checkdb rc 기록. 격리 dir 모드는 스크래치 백업으로 PASS 사전 검증됨.", color=GRAY, size=9.5)

pb()
h("4. S3 (오브젝트 스토리지) 검증")
h2("4.1 사전조건")
table(["점검","내용"],
 [["aws CLI","송신 서버에 설치(aws --version)"],
  ["자격증명","aws configure 또는 IAM 역할. 대상 버킷에 PutObject/ListBucket 권한"],
  ["버킷/경로","대상 버킷 존재. 업로드 prefix 결정(s3://<bucket>/cubrid/<db>)"],
  ["S3 호환(옵션)","MinIO/Ceph 등이면 AWS_ENDPOINT 지정, 필요시 AWS_PROFILE"],
  ["스테이징","로컬에 백업 크기만큼 여유"]],
 widths=[1.5,5.1])
h2("4.2 실행 명령")
code("DB=<db> LEVEL=0 STAGE=/backup/stage/<db> \\\n  S3_URI=s3://<bucket>/cubrid/<db> [AWS_PROFILE=<p>] [AWS_ENDPOINT=<url>] \\\n  bash s3_cubrid_backup.sh")
h2("4.3 기대결과 / 확인")
chk("스크립트 rc=0, 로그에 '[완료] ... S3 백업 성공'")
chk("aws s3 ls s3://<bucket>/cubrid/<db>/<TS>/ 에 <db>_bk0v000 존재(크기>0)")
chk("bpbackup 재시도 없이(또는 재시도 후) 업로드 성공")
h2("4.4 복원 검증")
code("aws s3 cp s3://<bucket>/cubrid/<db>/<TS>/ ./rst/ --recursive\ncubrid restoredb -B ./rst <스크래치DB> && cubrid checkdb --SA-mode <스크래치DB>")

h("5. NFS / CIFS NAS 검증")
h2("5.1 사전조건")
table(["점검","내용"],
 [["공유/Export","NAS 에서 NFS export 또는 CIFS 공유 준비, 쓰기 권한"],
  ["마운트 권한","클라이언트에서 mount 가능(root 또는 sudo). CIFS 는 자격증명(MOUNT_OPT)"],
  ["방화벽","NFS(2049 등)/CIFS(445) 통신 허용"],
  ["모드 결정","MODE=direct(간단) 또는 MODE=stage(운영 안전, 로컬 스테이징)"]],
 widths=[1.5,5.1])
h2("5.2 실행 명령")
code("DB=<db> MOUNT_SRC=<NAS>:/vol/backup MOUNTPOINT=/mnt/nasbackup FS_TYPE=nfs \\\n  MODE=direct DEST_SUBDIR=cubrid/<db> [MOUNT_OPT=<opts>] \\\n  bash nfs_cubrid_backup.sh")
h2("5.3 기대결과 / 확인")
chk("마운트 성공(mountpoint -q /mnt/nasbackup)")
chk("스크립트 rc=0, DEST(/mnt/nasbackup/cubrid/<db>/<TS>)에 <db>_bk0v000 존재")
chk("MODE=stage 인 경우 로컬 스테이징 -> NAS 복사 정상")
h2("5.4 복원 검증")
code("cubrid restoredb -B /mnt/nasbackup/cubrid/<db>/<TS> <스크래치DB>\ncubrid checkdb --SA-mode <스크래치DB>")

pb()
h("6. Veritas NetBackup 검증")
h2("6.1 사전조건")
table(["점검","내용"],
 [["NBU 클라이언트","DB 서버가 NetBackup 클라이언트로 등록됨"],
  ["정책/스케줄","마스터에 정책(POLICY)/스케줄(SCHEDULE) 구성(정책 구성 가이드 참조)"],
  ["연결 확인","bptestbpcd 또는 bpclntcmd -pn 성공"],
  ["경로/스테이징","bpbackup 경로(/usr/openv/netbackup/bin), 로컬 스테이징 공간"]],
 widths=[1.6,5.0])
h2("6.2 실행 명령")
code("DB=<db> LEVEL=0 STAGE=/backup/stage/<db> NBU_BIN=/usr/openv/netbackup/bin \\\n  POLICY=<정책명> SCHEDULE=<스케줄명> KEYWORD=cubrid_<db> \\\n  bash netbackup_cubrid_backup.sh")
h2("6.3 기대결과 / 확인")
chk("스크립트 rc=0, 로그에 'NetBackup 수집 성공' 및 '[완료]'")
chk("bplist 로 백업 이미지 확인(아래)")
chk("(순단 시) bpbackup 재시도 후 성공")
h2("6.4 복원 검증")
code("bplist -C <client> -t 0 -R /backup/stage/<db>\nbprestore -C <client> -D <client> -t 0 -w -L /tmp/rest.log /backup/stage/<db>\ncubrid restoredb -B <복원된경로> <스크래치DB> && cubrid checkdb --SA-mode <스크래치DB>")

h("7. 판정 기준 요약")
table(["방식","통과 조건"],
 [["S3","스크립트 rc=0 + s3 ls 파일 존재 + restoredb 0 + checkdb 0"],
  ["NFS/CIFS","마운트 성공 + rc=0 + NAS 파일 존재 + restoredb 0 + checkdb 0"],
  ["NetBackup","rc=0 + bplist 이미지 존재 + bprestore 0 + restoredb 0 + checkdb 0"]],
 widths=[1.4,5.2])

h("8. 정리(검증 후)")
chk("스크래치 DB 삭제(cubrid deletedb <스크래치DB>)")
chk("임시 복원/다운로드 파일 삭제")
chk("검증용 임시 자격증명/SSH 키 회수(불필요 시)")
chk("검증 결과와 소요시간/전송크기 기록")

h("9. 참고 - 격리(mock) 종단 시험과의 관계")
para("격리 시험 e2e_devices_test.sh 는 대상을 mock 으로 대체해 스크립트 로직과 CUBRID 백업/복원 무결성을 자동 검증한다(현재 4/4 PASS). 본 체크리스트는 그 스크립트를 mock 없이 실제 S3/NAS/NetBackup 에서 1회 검증하는 절차다. 실환경에서는 mock 대신 실제 aws/마운트/NBU 클라이언트만 준비하면 동일 스크립트를 그대로 사용하며, 검증은 realrun_verify.sh(3.1)로 자동화할 수 있다.")
para("파일 요약: 백업 스크립트(s3_/nfs_/netbackup_cubrid_backup.sh) · 격리 e2e(e2e_devices_test.sh) · 실환경 자동검증(realrun_verify.sh) · 본 체크리스트/매뉴얼/정책구성가이드(.docx)", color=GRAY, size=9.5)

doc.save(OUT); print("saved:", OUT)
